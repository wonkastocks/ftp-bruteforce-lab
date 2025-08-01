import streamlit as st
import datetime
import base64
from io import BytesIO
import pandas as pd
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
import openpyxl
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import random

# LinkedIn breach credentials - 10 accounts
USERS = {
    "michael123": "123456",
    "jennifer88": "password",
    "david2012": "12345678",
    "admin": "qwerty",
    "robert1": "abc123",
    "jessica": "password1",
    "john.smith": "123456789",
    "mary_jones": "letmein",
    "william": "monkey",
    "susan2023": "1234567"
}

def init_session_state():
    """Initialize session state variables"""
    if 'logged_in' not in st.session_state:
        st.session_state.logged_in = False
    if 'username' not in st.session_state:
        st.session_state.username = None
    if 'current_path' not in st.session_state:
        st.session_state.current_path = "/"
    if 'login_attempts' not in st.session_state:
        st.session_state.login_attempts = []
    if 'activity_logs' not in st.session_state:
        st.session_state.activity_logs = []

def log_attempt(username, password, success):
    """Log login attempts"""
    attempt = {
        "timestamp": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "username": username,
        "password": password,
        "success": success,
        "ip_address": "Client IP"  # In real app, would get actual IP
    }
    st.session_state.login_attempts.append(attempt)
    
def log_activity(username, action, details=""):
    """Log user activities"""
    activity = {
        "timestamp": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "username": username,
        "action": action,
        "details": details,
        "path": st.session_state.current_path
    }
    st.session_state.activity_logs.append(activity)

def create_sales_excel():
    """Create Excel file with fake sales data"""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Q3 2023 Sales"
    
    # Headers
    headers = ["Date", "Product", "Quantity", "Unit Price", "Total", "Region", "Sales Rep"]
    for col, header in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=header)
        ws.cell(row=1, column=col).font = openpyxl.styles.Font(bold=True)
    
    # Sample data
    products = ["Widget Pro", "Super Widget", "Widget Lite", "Enterprise Widget", "Widget Plus"]
    regions = ["North", "South", "East", "West", "Central"]
    reps = ["J. Smith", "M. Johnson", "S. Williams", "R. Brown", "K. Davis"]
    
    row = 2
    for month in range(7, 10):  # July to September
        for day in range(1, 29, 3):
            date = f"2023-{month:02d}-{day:02d}"
            product = random.choice(products)
            quantity = random.randint(10, 100)
            unit_price = random.randint(50, 500)
            total = quantity * unit_price
            region = random.choice(regions)
            rep = random.choice(reps)
            
            ws.cell(row=row, column=1, value=date)
            ws.cell(row=row, column=2, value=product)
            ws.cell(row=row, column=3, value=quantity)
            ws.cell(row=row, column=4, value=f"${unit_price}")
            ws.cell(row=row, column=5, value=f"${total:,}")
            ws.cell(row=row, column=6, value=region)
            ws.cell(row=row, column=7, value=rep)
            row += 1
    
    # Auto-adjust column widths
    for column in ws.columns:
        max_length = max(len(str(cell.value)) for cell in column if cell.value)
        ws.column_dimensions[column[0].column_letter].width = max_length + 2
    
    # Save to bytes
    excel_buffer = BytesIO()
    wb.save(excel_buffer)
    excel_buffer.seek(0)
    return excel_buffer.getvalue()

def create_financial_pdf(quarter, revenue, expenses):
    """Create PDF financial report"""
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    
    # Header with ISG branding
    c.setFont("Helvetica-Bold", 24)
    c.drawString(50, height - 50, "ISG Cybersecurity")
    c.setFont("Helvetica", 14)
    c.drawString(50, height - 75, "There's No Challenge Too Big")
    c.setFont("Helvetica-Bold", 18)
    c.drawString(50, height - 110, f"Security Services Report - {quarter}")
    
    # Line
    c.line(50, height - 130, width - 50, height - 130)
    
    # Financial Summary
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, height - 160, "Executive Summary")
    c.setFont("Helvetica", 12)
    c.drawString(50, height - 190, f"Total Revenue: ${revenue}")
    c.drawString(50, height - 210, f"Total Expenses: ${expenses}")
    c.drawString(50, height - 230, f"Net Profit: ${int(revenue.replace('M', '')) - int(expenses.replace('M', ''))}M")
    
    # Details
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, height - 280, "Revenue Breakdown")
    c.setFont("Helvetica", 11)
    y_pos = height - 310
    categories = [
        ("Security Assessments", "40%"),
        ("Incident Response", "25%"),
        ("Compliance Audits", "20%"),
        ("Security Training", "15%")
    ]
    for category, percentage in categories:
        c.drawString(70, y_pos, f"• {category}: {percentage}")
        y_pos -= 20
    
    # Footer
    c.setFont("Helvetica", 10)
    c.drawString(50, 50, "Confidential - ISG Cybersecurity")
    c.drawString(width - 150, 50, f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d')}")
    
    c.save()
    buffer.seek(0)
    return buffer.getvalue()

def create_employee_doc():
    """Create Word document with employee handbook"""
    doc = Document()
    
    # Title
    title = doc.add_heading('ISG Cybersecurity Employee Handbook', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Tagline
    tagline = doc.add_paragraph("There's No Challenge Too Big")
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add company info
    doc.add_heading('Welcome to ISG Cybersecurity', level=1)
    doc.add_paragraph(
        'ISG Cybersecurity is a leading cybersecurity firm specializing in '
        'protecting organizations from digital threats. We guide our clients '
        'through the mountain of cyber challenges with expertise and dedication.'
    )
    
    # Mission Statement
    doc.add_heading('Our Mission', level=2)
    doc.add_paragraph(
        'To guide organizations safely through the evolving landscape of cyber threats, '
        'providing expert protection and ensuring that no cybersecurity challenge is too big to overcome.'
    )
    
    # Company Policies
    doc.add_heading('Company Policies', level=1)
    doc.add_heading('Work Hours', level=2)
    doc.add_paragraph('Standard work hours are 9:00 AM to 5:00 PM, Monday through Friday.')
    
    doc.add_heading('Remote Work Policy', level=2)
    doc.add_paragraph(
        'Employees may work remotely up to 3 days per week with manager approval. '
        'All remote work must be conducted from a secure location with reliable internet.'
    )
    
    # Benefits
    doc.add_heading('Employee Benefits', level=1)
    benefits = [
        'Health, Dental, and Vision Insurance',
        '401(k) with company matching',
        '20 days PTO + holidays',
        'Professional development budget',
        'Gym membership reimbursement'
    ]
    for benefit in benefits:
        doc.add_paragraph(f'• {benefit}', style='List Bullet')
    
    # Save to bytes
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer.getvalue()

def create_database_backup():
    """Create SQL backup file"""
    sql_content = """-- ISG Cybersecurity Database Backup
-- Generated: 2023-12-01 03:00:00
-- Database: isg_cyber_prod

SET SQL_MODE = "NO_AUTO_VALUE_ON_ZERO";
START TRANSACTION;
SET time_zone = "+00:00";

-- Database: `isg_cyber_prod`
CREATE DATABASE IF NOT EXISTS `isg_cyber_prod` DEFAULT CHARACTER SET utf8mb4;
USE `isg_cyber_prod`;

-- Table structure for table `employees`
DROP TABLE IF EXISTS `employees`;
CREATE TABLE `employees` (
  `id` int(11) NOT NULL AUTO_INCREMENT,
  `employee_id` varchar(10) NOT NULL,
  `first_name` varchar(50) NOT NULL,
  `last_name` varchar(50) NOT NULL,
  `email` varchar(100) NOT NULL,
  `department` varchar(50) NOT NULL,
  `hire_date` date NOT NULL,
  `salary` decimal(10,2) NOT NULL,
  PRIMARY KEY (`id`),
  UNIQUE KEY `employee_id` (`employee_id`)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;

-- Sample data for `employees`
INSERT INTO `employees` VALUES
(1,'EMP001','John','Doe','john.doe@isgcyber.com','Security Analysis','2020-01-15',95000.00),
(2,'EMP002','Jane','Smith','jane.smith@isgcyber.com','Incident Response','2019-03-22',85000.00),
(3,'EMP003','Michael','Johnson','michael.j@isgcyber.com','Compliance','2021-06-10',78000.00),
(4,'EMP004','Sarah','Williams','sarah.w@isgcyber.com','HR','2018-11-05',72000.00),
(5,'EMP005','Robert','Brown','robert.b@isgcyber.com','Penetration Testing','2022-02-28',102000.00);

-- Table structure for table `customers`
DROP TABLE IF EXISTS `customers`;
CREATE TABLE `customers` (
  `id` int(11) NOT NULL AUTO_INCREMENT,
  `company_name` varchar(100) NOT NULL,
  `contact_name` varchar(100) NOT NULL,
  `email` varchar(100) NOT NULL,
  `phone` varchar(20) NOT NULL,
  `contract_value` decimal(12,2) NOT NULL,
  PRIMARY KEY (`id`)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;

COMMIT;
"""
    return sql_content.encode()

def login_page():
    """Display login page"""
    st.markdown("""
    <style>
    .login-container {
        background-color: #f0f0f0;
        padding: 20px;
        border-radius: 5px;
        margin: 20px auto;
        max-width: 400px;
    }
    .isg-header {
        background-color: #1e467e;
        color: white;
        padding: 20px;
        text-align: center;
        margin: -10px -10px 20px -10px;
    }
    .isg-title {
        font-size: 36px;
        font-weight: bold;
        margin: 0;
    }
    .isg-tagline {
        font-size: 16px;
        font-style: italic;
        margin-top: 5px;
    }
    </style>
    """, unsafe_allow_html=True)
    
    st.markdown("""
    <div class="isg-header">
        <div class="isg-title">ISG Cybersecurity</div>
        <div class="isg-tagline">There's No Challenge Too Big</div>
    </div>
    """, unsafe_allow_html=True)
    
    st.title("🏔️ FTP Server Portal")
    st.markdown("---")
    
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        st.markdown("### FTP Login")
        
        with st.form("login_form"):
            username = st.text_input("Username", placeholder="Enter username")
            password = st.text_input("Password", type="password", placeholder="Enter password")
            login_button = st.form_submit_button("Login", use_container_width=True)
            
            if login_button:
                if username in USERS and USERS[username] == password:
                    st.session_state.logged_in = True
                    st.session_state.username = username
                    log_attempt(username, password, True)
                    log_activity(username, "LOGIN", f"Successful login from web interface")
                    st.success("Login successful!")
                    st.rerun()
                else:
                    log_attempt(username, password, False)
                    st.error("Invalid username or password")
        
        st.markdown("---")
        st.caption("ISG Cybersecurity FTP Server v2.1.3 | Port: 21 | Protocol: FTP/FTPS")

def file_browser():
    """Display file browser interface with traditional FTP listing"""
    # Custom CSS for FTP-like interface with ISG branding
    st.markdown("""
    <style>
    .ftp-listing {
        font-family: monospace;
        background-color: #1e1e1e;
        color: #00ff00;
        padding: 10px;
        border-radius: 5px;
        margin: 10px 0;
    }
    .ftp-header {
        color: #ffffff;
        background-color: #1e467e;
        padding: 5px 10px;
        margin-bottom: 10px;
        border-radius: 3px;
    }
    .stButton > button {
        background-color: transparent;
        color: #1e467e;
        border: none;
        padding: 0;
        text-decoration: underline;
        font-family: monospace;
    }
    .isg-banner {
        background-color: #1e467e;
        color: white;
        padding: 10px;
        border-radius: 5px;
        margin-bottom: 20px;
        text-align: center;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # ISG Banner
    st.markdown("""
    <div class="isg-banner">
        <strong>ISG Cybersecurity</strong> | Secure FTP Portal | There's No Challenge Too Big
    </div>
    """, unsafe_allow_html=True)
    
    # Header
    col1, col2 = st.columns([6, 1])
    with col1:
        st.markdown(f"### 🏔️ FTP File Browser - Connected as: **{st.session_state.username}**")
    with col2:
        if st.button("Logout", type="secondary"):
            st.session_state.logged_in = False
            st.session_state.username = None
            st.session_state.current_path = "/"
            st.rerun()
    
    # Current directory display
    st.markdown(f"**Current Directory:** `{st.session_state.current_path}`")
    
    # File listing header
    st.markdown("""
    <div class="ftp-header">
    <pre>Permissions    Size     Date Modified    Name</pre>
    </div>
    """, unsafe_allow_html=True)
    
    # Navigation
    if st.session_state.current_path != "/":
        if st.button("📁 ..", key="parent_dir"):
            st.session_state.current_path = "/"
            st.rerun()
    
    # Define file structure with real files
    if st.session_state.current_path == "/":
        files = [
            ("drwxr-xr-x", "4096", "Nov 15 09:30", "public", True),
            ("drwx------", "4096", "Nov 20 14:15", "private", True),
            ("drwxr-xr-x", "4096", "Dec 01 08:45", "reports", True),
            ("drwxr-xr-x", "4096", "Dec 01 03:00", "backups", True),
        ]
        # Add LOG folder only for admin
        if st.session_state.username == "admin":
            files.append(("drwx------", "4096", datetime.datetime.now().strftime("%b %d %H:%M"), "LOG", True))
    elif st.session_state.current_path == "/public":
        files = [
            ("-rw-r--r--", "1234", "Nov 10 12:00", "readme.txt", False),
            ("-rw-r--r--", "45678", "Nov 12 15:30", "company_brochure.pdf", False),
            ("-rw-r--r--", "23456", "Nov 14 09:15", "welcome_packet.docx", False),
            ("-rw-r--r--", "34567", "Dec 18 14:20", "5_ways_secure_business.pdf", False),
            ("-rw-r--r--", "28901", "Dec 19 10:15", "cybersecurity_assessment_guide.pdf", False),
        ]
    elif st.session_state.current_path == "/private":
        files = [
            ("-rw-------", "98765", "Nov 18 16:45", "salary_data.xlsx", False),
            ("-rw-------", "2345", "Nov 19 10:30", "passwords.txt", False),
            ("-rw-------", "87654", "Nov 20 14:00", "employee_records.xlsx", False),
        ]
    elif st.session_state.current_path == "/reports":
        files = [
            ("-rw-r--r--", "156789", "Sep 30 17:00", "Q3_2023_financial.pdf", False),
            ("-rw-r--r--", "167890", "Oct 31 17:00", "Q4_2023_financial.pdf", False),
            ("-rw-r--r--", "234567", "Dec 15 12:00", "annual_report_2023.docx", False),
            ("-rw-r--r--", "345678", "Dec 20 14:30", "sales_analysis_2023.xlsx", False),
            ("-rw-r--r--", "456789", "Dec 22 11:45", "penetration_test_report.pdf", False),
            ("-rw-r--r--", "234890", "Dec 23 09:30", "incident_response_summary.docx", False),
        ]
    elif st.session_state.current_path == "/backups":
        files = [
            ("-rw-r--r--", "5678901", "Dec 01 03:00", "database_backup_20231201.sql", False),
            ("-rw-r--r--", "8901234", "Dec 01 03:00", "website_backup.tar.gz", False),
            ("-rw-r--r--", "2345678", "Dec 01 03:00", "config_backup.zip", False),
        ]
    elif st.session_state.current_path == "/LOG" and st.session_state.username == "admin":
        # Admin-only LOG folder
        files = [
            ("-rw-------", str(len(str(st.session_state.login_attempts))), 
             datetime.datetime.now().strftime("%b %d %H:%M"), "login_attempts.log", False),
            ("-rw-------", str(len(str(st.session_state.activity_logs))), 
             datetime.datetime.now().strftime("%b %d %H:%M"), "user_activity.log", False),
            ("-rw-------", "2048", datetime.datetime.now().strftime("%b %d %H:%M"), "access_summary.txt", False),
            ("-rw-------", "4096", datetime.datetime.now().strftime("%b %d %H:%M"), "security_audit.log", False),
        ]
    else:
        files = []
    
    # Display files
    for perms, size, date, name, is_dir in files:
        col1, col2, col3, col4 = st.columns([3, 1, 2, 4])
        
        with col1:
            st.text(perms)
        with col2:
            st.text(size)
        with col3:
            st.text(date)
        with col4:
            if is_dir:
                if st.button(f"📁 {name}", key=f"dir_{name}"):
                    st.session_state.current_path = f"/{name}"
                    log_activity(st.session_state.username, "NAVIGATE", f"Changed directory to {name}")
                    st.rerun()
            else:
                # Generate file content based on filename
                file_content = generate_file_content(name)
                
                download_button = st.download_button(
                    label=f"📄 {name}",
                    data=file_content,
                    file_name=name,
                    mime=get_mime_type(name),
                    key=f"download_{name}",
                    on_click=lambda n=name: log_activity(st.session_state.username, "DOWNLOAD", f"Downloaded file: {n}")
                )
    
    # Server info
    st.markdown("---")
    st.caption(f"ISG Cybersecurity FTP Server | ftp.isgcyber.com | Connected: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | Transfer Mode: Binary")

def create_isg_checklist_pdf():
    """Create ISG's 5 Ways to Secure Your Business checklist"""
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    
    # Header
    c.setFont("Helvetica-Bold", 24)
    c.drawString(50, height - 50, "ISG Cybersecurity")
    c.setFont("Helvetica", 14)
    c.drawString(50, height - 75, "There's No Challenge Too Big")
    
    # Title
    c.setFont("Helvetica-Bold", 20)
    c.drawString(50, height - 120, "5 Simple Ways to Secure Your Business")
    c.line(50, height - 130, width - 50, height - 130)
    
    # Checklist items
    c.setFont("Helvetica-Bold", 14)
    y_pos = height - 170
    
    checklist = [
        ("1. Enable Multi-Factor Authentication (MFA)",
         "Require MFA for all user accounts, especially admin and privileged users."),
        ("2. Keep Software Updated", 
         "Install security patches promptly and enable automatic updates where possible."),
        ("3. Train Employees on Cybersecurity",
         "Conduct regular security awareness training to recognize phishing and social engineering."),
        ("4. Implement Strong Password Policies",
         "Require complex passwords and regular password changes. Use a password manager."),
        ("5. Regular Backups and Testing",
         "Backup critical data regularly and test restore procedures quarterly.")
    ]
    
    for title, desc in checklist:
        c.setFont("Helvetica-Bold", 12)
        c.drawString(70, y_pos, title)
        c.setFont("Helvetica", 11)
        c.drawString(90, y_pos - 20, desc)
        y_pos -= 60
    
    # Footer
    c.setFont("Helvetica", 10)
    c.drawString(50, 50, "© ISG Cybersecurity | www.isgcyber.com")
    
    c.save()
    buffer.seek(0)
    return buffer.getvalue()

def create_assessment_guide_pdf():
    """Create Cybersecurity Assessment Guide PDF"""
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    
    # Header
    c.setFont("Helvetica-Bold", 24)
    c.drawString(50, height - 50, "ISG Cybersecurity")
    c.setFont("Helvetica-Bold", 18)
    c.drawString(50, height - 90, "Cybersecurity Assessment Guide")
    c.line(50, height - 100, width - 50, height - 100)
    
    # Content
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, height - 130, "Assessment Framework")
    
    c.setFont("Helvetica", 11)
    y_pos = height - 160
    
    sections = [
        "Phase 1: Discovery and Reconnaissance",
        "• Network topology mapping",
        "• Asset identification",
        "• Service enumeration",
        "",
        "Phase 2: Vulnerability Analysis",
        "• Automated scanning",
        "• Manual verification",
        "• Risk classification",
        "",
        "Phase 3: Exploitation Testing",
        "• Controlled penetration testing",
        "• Privilege escalation attempts",
        "• Lateral movement analysis",
        "",
        "Phase 4: Reporting and Remediation",
        "• Executive summary",
        "• Technical findings",
        "• Remediation roadmap"
    ]
    
    for section in sections:
        c.drawString(70, y_pos, section)
        y_pos -= 20
    
    c.save()
    buffer.seek(0)
    return buffer.getvalue()

def create_pentest_report_pdf():
    """Create Penetration Test Report PDF"""
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    
    # Header
    c.setFont("Helvetica-Bold", 24)
    c.drawString(50, height - 50, "ISG Cybersecurity")
    c.setFont("Helvetica-Bold", 18)
    c.drawString(50, height - 90, "Penetration Test Report - Executive Summary")
    c.line(50, height - 100, width - 50, height - 100)
    
    # Client info
    c.setFont("Helvetica", 12)
    c.drawString(50, height - 130, "Client: Confidential Corp")
    c.drawString(50, height - 150, "Test Period: December 1-15, 2023")
    c.drawString(50, height - 170, "Test Type: External Network Penetration Test")
    
    # Findings summary
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, height - 210, "Key Findings")
    
    c.setFont("Helvetica", 11)
    y_pos = height - 240
    
    findings = [
        "Critical: 2 findings requiring immediate attention",
        "High: 5 findings requiring prompt remediation",
        "Medium: 8 findings requiring scheduled fixes",
        "Low: 12 informational findings"
    ]
    
    for finding in findings:
        c.drawString(70, y_pos, f"• {finding}")
        y_pos -= 20
    
    # Risk score
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y_pos - 20, "Overall Risk Score: 7.2/10 (High)")
    
    c.save()
    buffer.seek(0)
    return buffer.getvalue()

def create_isg_company_brochure_pdf():
    """Create ISG Cybersecurity Company Brochure"""
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    
    # Header
    c.setFont("Helvetica-Bold", 28)
    c.drawString(50, height - 60, "ISG Cybersecurity")
    c.setFont("Helvetica-Oblique", 16)
    c.drawString(50, height - 85, "There's No Challenge Too Big")
    
    # Mountain graphic (simple representation)
    c.setLineWidth(2)
    c.line(100, height - 120, 150, height - 170)
    c.line(150, height - 170, 200, height - 120)
    
    # Services section
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, height - 220, "Our Services")
    c.line(50, height - 225, 200, height - 225)
    
    c.setFont("Helvetica", 12)
    y_pos = height - 250
    
    services = [
        "• Cybersecurity Assessments - Comprehensive security evaluations",
        "• Penetration Testing - Real-world attack simulations",
        "• Incident Response - 24/7 emergency response team",
        "• Security Operations Center - Continuous monitoring",
        "• Compliance Audits - SOC2, PCI-DSS, HIPAA",
        "• Security Training - Employee awareness programs"
    ]
    
    for service in services:
        c.drawString(70, y_pos, service)
        y_pos -= 25
    
    # Why ISG section
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, y_pos - 30, "Why Choose ISG?")
    c.line(50, y_pos - 35, 180, y_pos - 35)
    
    c.setFont("Helvetica", 12)
    y_pos = y_pos - 60
    
    reasons = [
        "✓ Experienced team of certified security professionals",
        "✓ Proven track record protecting Fortune 500 companies",
        "✓ Cutting-edge tools and methodologies",
        "✓ Tailored solutions for your unique challenges"
    ]
    
    for reason in reasons:
        c.drawString(70, y_pos, reason)
        y_pos -= 20
    
    # Footer
    c.setFont("Helvetica", 10)
    c.drawString(50, 50, "Contact: info@isgcyber.com | www.isgcyber.com")
    
    c.save()
    buffer.seek(0)
    return buffer.getvalue()

def create_incident_response_doc():
    """Create Incident Response Summary Document"""
    doc = Document()
    
    # Title
    title = doc.add_heading('ISG Cybersecurity', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Incident Response Summary', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Incident details
    doc.add_heading('Incident Overview', 2)
    doc.add_paragraph('Date: December 23, 2023')
    doc.add_paragraph('Incident Type: Ransomware Attack (Contained)')
    doc.add_paragraph('Severity: High')
    doc.add_paragraph('Status: Resolved')
    
    # Timeline
    doc.add_heading('Response Timeline', 2)
    timeline = [
        '08:45 - Initial detection by SOC team',
        '09:00 - Incident response team activated',
        '09:30 - Affected systems isolated',
        '10:15 - Malware identified as BlackCat variant',
        '11:00 - Containment measures implemented',
        '14:00 - Systems restored from clean backups',
        '16:30 - Normal operations resumed'
    ]
    
    for event in timeline:
        doc.add_paragraph(event, style='List Bullet')
    
    # Recommendations
    doc.add_heading('Post-Incident Recommendations', 2)
    doc.add_paragraph(
        'Based on this incident, ISG Cybersecurity recommends implementing '
        'enhanced endpoint detection and response (EDR) solutions and '
        'conducting quarterly tabletop exercises.'
    )
    
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer.getvalue()

def create_login_attempts_log():
    """Create login attempts log file"""
    log_content = "ISG Cybersecurity FTP Server - Login Attempts Log\n"
    log_content += "=" * 60 + "\n"
    log_content += f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
    log_content += "=" * 60 + "\n\n"
    
    for attempt in st.session_state.login_attempts:
        status = "SUCCESS" if attempt['success'] else "FAILED"
        log_content += f"{attempt['timestamp']} | {status} | Username: {attempt['username']} | Password: {attempt['password']} | IP: {attempt['ip_address']}\n"
    
    return log_content.encode()

def create_activity_log():
    """Create user activity log file"""
    log_content = "ISG Cybersecurity FTP Server - User Activity Log\n"
    log_content += "=" * 60 + "\n"
    log_content += f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
    log_content += "=" * 60 + "\n\n"
    
    for activity in st.session_state.activity_logs:
        log_content += f"{activity['timestamp']} | User: {activity['username']} | Action: {activity['action']} | Path: {activity['path']} | Details: {activity['details']}\n"
    
    return log_content.encode()

def create_access_summary():
    """Create access summary report"""
    summary = "ISG Cybersecurity FTP Server - Access Summary Report\n"
    summary += "=" * 60 + "\n"
    summary += f"Report Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
    summary += "=" * 60 + "\n\n"
    
    # Login statistics
    total_attempts = len(st.session_state.login_attempts)
    successful = len([a for a in st.session_state.login_attempts if a['success']])
    failed = total_attempts - successful
    
    summary += "LOGIN STATISTICS:\n"
    summary += f"Total Login Attempts: {total_attempts}\n"
    summary += f"Successful Logins: {successful}\n"
    summary += f"Failed Attempts: {failed}\n"
    summary += f"Success Rate: {(successful/total_attempts*100 if total_attempts > 0 else 0):.1f}%\n\n"
    
    # User activity statistics
    summary += "USER ACTIVITY SUMMARY:\n"
    activity_counts = {}
    for activity in st.session_state.activity_logs:
        action = activity['action']
        activity_counts[action] = activity_counts.get(action, 0) + 1
    
    for action, count in activity_counts.items():
        summary += f"{action}: {count} times\n"
    
    # Most active users
    user_actions = {}
    for activity in st.session_state.activity_logs:
        user = activity['username']
        user_actions[user] = user_actions.get(user, 0) + 1
    
    summary += "\nMOST ACTIVE USERS:\n"
    for user, count in sorted(user_actions.items(), key=lambda x: x[1], reverse=True)[:5]:
        summary += f"{user}: {count} actions\n"
    
    return summary.encode()

def create_security_audit_log():
    """Create security audit log"""
    audit = "ISG Cybersecurity FTP Server - Security Audit Log\n"
    audit += "=" * 60 + "\n"
    audit += f"Audit Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
    audit += "=" * 60 + "\n\n"
    
    # Failed login analysis
    audit += "FAILED LOGIN ANALYSIS:\n"
    failed_by_user = {}
    for attempt in st.session_state.login_attempts:
        if not attempt['success']:
            user = attempt['username']
            failed_by_user[user] = failed_by_user.get(user, 0) + 1
    
    audit += "Users with failed login attempts:\n"
    for user, count in sorted(failed_by_user.items(), key=lambda x: x[1], reverse=True):
        audit += f"  {user}: {count} failed attempts\n"
    
    # Suspicious activity
    audit += "\nSUSPICIOUS ACTIVITY FLAGS:\n"
    for user, count in failed_by_user.items():
        if count >= 3:
            audit += f"⚠️  WARNING: {user} has {count} failed login attempts\n"
    
    # File access patterns
    audit += "\nFILE ACCESS PATTERNS:\n"
    downloads = [a for a in st.session_state.activity_logs if a['action'] == 'DOWNLOAD']
    if downloads:
        audit += f"Total files downloaded: {len(downloads)}\n"
        for d in downloads[-10:]:  # Last 10 downloads
            audit += f"  {d['timestamp']} - {d['username']} downloaded {d['details']}\n"
    
    # Recommendations
    audit += "\nSECURITY RECOMMENDATIONS:\n"
    if any(count >= 5 for count in failed_by_user.values()):
        audit += "- Implement account lockout after 5 failed attempts\n"
    total_attempts = len(st.session_state.login_attempts)
    if total_attempts > 100:
        audit += "- High volume of login attempts detected - consider rate limiting\n"
    audit += "- Enable two-factor authentication for all users\n"
    audit += "- Regularly review and rotate passwords\n"
    
    return audit.encode()

def generate_file_content(filename):
    """Generate appropriate content based on filename"""
    # Handle LOG folder files for admin
    if filename == "login_attempts.log":
        return create_login_attempts_log()
    elif filename == "user_activity.log":
        return create_activity_log()
    elif filename == "access_summary.txt":
        return create_access_summary()
    elif filename == "security_audit.log":
        return create_security_audit_log()
    elif filename == "readme.txt":
        return b"""ISG Cybersecurity FTP Server
=============================
There's No Challenge Too Big

Welcome to the ISG Cybersecurity FTP server. This server contains security
assessments, reports, and resources for authorized personnel only.

Directory Structure:
- /public    : Public security advisories and resources
- /private   : Confidential client assessments (restricted access)
- /reports   : Security audit reports and compliance documents
- /backups   : System and security configuration backups

For support, contact: support@isgcyber.com
Website: https://isgcyber.com
"""
    
    elif filename == "passwords.txt":
        return b"""CONFIDENTIAL - System Passwords
================================

Database Server:
- Host: db.securecorp.local
- Admin: dbadmin / P@ssw0rd2023!

Email Server:
- Host: mail.securecorp.com
- Admin: mailadmin / M@ilSecure#456

Backup System:
- Host: backup.securecorp.local
- Admin: backupadmin / B@ckup!789

WiFi Networks:
- SecureCorp-Guest: Welcome2023
- SecureCorp-Staff: Staff#Secure$2023

IMPORTANT: Change these passwords regularly!
"""
    
    elif filename.endswith('.pdf'):
        if 'financial' in filename:
            if 'Q3' in filename:
                return create_financial_pdf("Q3 2023", "2.5M", "1.8M")
            elif 'Q4' in filename:
                return create_financial_pdf("Q4 2023", "3.1M", "2.0M")
            else:
                return create_financial_pdf("Q2 2023", "2.2M", "1.6M")
        elif '5_ways_secure_business' in filename:
            return create_isg_checklist_pdf()
        elif 'cybersecurity_assessment_guide' in filename:
            return create_assessment_guide_pdf()
        elif 'penetration_test_report' in filename:
            return create_pentest_report_pdf()
        else:
            # Company brochure
            return create_isg_company_brochure_pdf()
    
    elif filename.endswith('.xlsx'):
        if 'salary' in filename or 'employee' in filename:
            return create_employee_excel()
        else:
            return create_sales_excel()
    
    elif filename.endswith('.docx'):
        if 'incident_response' in filename:
            return create_incident_response_doc()
        else:
            return create_employee_doc()
    
    elif filename.endswith('.sql'):
        return create_database_backup()
    
    elif filename.endswith('.zip') or filename.endswith('.tar.gz'):
        # Create a simple text file as placeholder
        return b"[Binary archive file - Contains compressed backup data]"
    
    else:
        return b"File content not available"

def create_employee_excel():
    """Create Excel with employee data"""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Employee Data"
    
    # Headers
    headers = ["Employee ID", "Name", "Department", "Position", "Salary", "Start Date", "Email"]
    for col, header in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=header)
        ws.cell(row=1, column=col).font = openpyxl.styles.Font(bold=True)
    
    # Employee data
    employees = [
        ("EMP001", "John Doe", "Security Analysis", "Senior Security Analyst", "$95,000", "2020-01-15", "john.doe@isgcyber.com"),
        ("EMP002", "Jane Smith", "Incident Response", "IR Team Lead", "$85,000", "2019-03-22", "jane.smith@isgcyber.com"),
        ("EMP003", "Michael Johnson", "Compliance", "Compliance Auditor", "$78,000", "2021-06-10", "michael.j@isgcyber.com"),
        ("EMP004", "Sarah Williams", "HR", "HR Specialist", "$72,000", "2018-11-05", "sarah.w@isgcyber.com"),
        ("EMP005", "Robert Brown", "Penetration Testing", "Senior Pentester", "$102,000", "2022-02-28", "robert.b@isgcyber.com"),
        ("EMP006", "Lisa Davis", "Threat Intelligence", "Threat Analyst", "$80,000", "2020-09-14", "lisa.d@isgcyber.com"),
        ("EMP007", "James Wilson", "Security Operations", "SOC Manager", "$88,000", "2019-07-20", "james.w@isgcyber.com"),
        ("EMP008", "Patricia Garcia", "Business Development", "Sales Director", "$110,000", "2017-04-03", "patricia.g@isgcyber.com"),
    ]
    
    for row, employee in enumerate(employees, 2):
        for col, value in enumerate(employee, 1):
            ws.cell(row=row, column=col, value=value)
    
    # Auto-adjust columns
    for column in ws.columns:
        max_length = max(len(str(cell.value)) for cell in column if cell.value)
        ws.column_dimensions[column[0].column_letter].width = max_length + 2
    
    excel_buffer = BytesIO()
    wb.save(excel_buffer)
    excel_buffer.seek(0)
    return excel_buffer.getvalue()

def get_mime_type(filename):
    """Get MIME type based on file extension"""
    if filename.endswith('.txt'):
        return 'text/plain'
    elif filename.endswith('.pdf'):
        return 'application/pdf'
    elif filename.endswith('.xlsx'):
        return 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    elif filename.endswith('.docx'):
        return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    elif filename.endswith('.sql'):
        return 'application/sql'
    elif filename.endswith('.zip'):
        return 'application/zip'
    elif filename.endswith('.tar.gz'):
        return 'application/gzip'
    else:
        return 'application/octet-stream'

def admin_panel():
    """Display admin panel with login attempts"""
    if st.session_state.username == "admin":
        st.markdown("---")
        st.markdown("### 🔍 Admin Panel - Login Attempts")
        
        if st.session_state.login_attempts:
            df = pd.DataFrame(st.session_state.login_attempts)
            st.dataframe(df, use_container_width=True)
            
            # Stats
            total_attempts = len(df)
            successful = len(df[df['success'] == True])
            failed = len(df[df['success'] == False])
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Total Attempts", total_attempts)
            with col2:
                st.metric("Successful", successful)
            with col3:
                st.metric("Failed", failed)
            
            # Download attempts as CSV
            csv = df.to_csv(index=False)
            st.download_button(
                label="📥 Download Login Log",
                data=csv,
                file_name=f"ftp_login_attempts_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv"
            )
        else:
            st.info("No login attempts recorded yet")

def main():
    """Main application"""
    st.set_page_config(
        page_title="ISG Cybersecurity FTP Server",
        page_icon="🏔️",
        layout="wide"
    )
    
    init_session_state()
    
    if not st.session_state.logged_in:
        login_page()
    else:
        file_browser()
        admin_panel()

if __name__ == "__main__":
    main()